"""
Document processing module for extracting and chunking text from various file formats
"""

import PyPDF2
from docx import Document
from typing import List, Tuple
import config


def extract_text_from_pdf(file) -> str:
    """
    Extract text from a PDF file
    
    Args:
        file: File object from Streamlit file uploader
        
    Returns:
        Extracted text as a string
    """
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"Error reading PDF file: {str(e)}")


def extract_text_from_txt(file) -> str:
    """
    Extract text from a TXT file
    
    Args:
        file: File object from Streamlit file uploader
        
    Returns:
        Extracted text as a string
    """
    try:
        text = file.read().decode('utf-8')
        return text.strip()
    except Exception as e:
        raise Exception(f"Error reading TXT file: {str(e)}")


def extract_text_from_docx(file) -> str:
    """
    Extract text from a DOCX file with enhanced extraction
    
    Args:
        file: File object from Streamlit file uploader
        
    Returns:
        Extracted text as a string
    """
    try:
        doc = Document(file)
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        text = "\n".join(text_parts)
        
        if not text.strip():
            raise ValueError("No text content found in DOCX file")
        
        return text
    except Exception as e:
        raise Exception(f"Error reading DOCX file: {str(e)}")


def extract_text(file, filename: str) -> str:
    """
    Extract text from a file based on its extension
    
    Args:
        file: File object from Streamlit file uploader
        filename: Name of the file
        
    Returns:
        Extracted text as a string
    """
    file_extension = filename.lower().split('.')[-1]
    
    if file_extension == 'pdf':
        return extract_text_from_pdf(file)
    elif file_extension == 'txt':
        return extract_text_from_txt(file)
    elif file_extension == 'docx':
        return extract_text_from_docx(file)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")


def chunk_text(text: str, filename: str, chunk_size: int = None, overlap: int = None) -> List[Tuple[str, dict]]:
    """
    Split text into overlapping chunks
    
    Args:
        text: Text to chunk
        filename: Original filename for metadata
        chunk_size: Size of each chunk in characters
        overlap: Overlap between chunks in characters
        
    Returns:
        List of tuples containing (chunk_text, metadata)
    """
    if chunk_size is None:
        chunk_size = config.CHUNK_SIZE
    if overlap is None:
        overlap = config.CHUNK_OVERLAP
    
    chunks = []
    start = 0
    chunk_index = 0
    
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        
        # Only add non-empty chunks
        if chunk.strip():
            metadata = {
                'filename': filename,
                'chunk_index': chunk_index,
                'total_chars': len(chunk)
            }
            chunks.append((chunk, metadata))
            chunk_index += 1
        
        start += chunk_size - overlap
    
    return chunks


def process_document(file, filename: str) -> List[Tuple[str, dict]]:
    """
    Process a document: extract text and chunk it
    
    Args:
        file: File object from Streamlit file uploader
        filename: Name of the file
        
    Returns:
        List of tuples containing (chunk_text, metadata)
    """
    text = extract_text(file, filename)
    
    if not text:
        raise ValueError(f"No text could be extracted from {filename}")
    
    chunks = chunk_text(text, filename)
    
    return chunks

